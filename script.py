import random
import re
import os
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx2pdf import convert

# --- SENTENCE SPLIT ---
def split_sentences(text):
    return re.split(r'(?<=[.!?]) +', text)

# --- HUMANIZE ---
def humanize_text(text):
    sentences = split_sentences(text)
    new_sentences = []

    for s in sentences:
        words = s.split()

        if len(words) > 18 and random.random() < 0.4:
            split = len(words) // 2
            new_sentences.append(" ".join(words[:split]) + ".")
            new_sentences.append(" ".join(words[split:]))
            continue

        if random.random() < 0.2:
            s = s.replace(" is ", random.choice([" is ", " is actually ", " is kind of "]))

        if random.random() < 0.15 and len(words) > 8:
            insert_pos = len(words) // 2
            words.insert(insert_pos, random.choice(["in a way", "to some extent"]))
            s = " ".join(words)

        new_sentences.append(s)

    return " ".join(new_sentences)

# --- EXTRA HUMANIZATION ---
def additional_humanization(text):
    sentences = split_sentences(text)

    for i in range(len(sentences)):
        if random.random() < 0.15 and sentences[i].endswith("."):
            sentences[i] = sentences[i][:-1] + ".."

    return " ".join(sentences)

# --- FORMAT ---
def format_paragraph(p):
    for run in p.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(12)

    p.paragraph_format.line_spacing = 2

# --- APA ---
def apply_apa(doc, text, name, instructor, course, date, title):
    section = doc.sections[0]

    header = section.header.paragraphs[0]
    header.text = f"Running head: {title.upper()}"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.add_run(f"{title}\n\n{name}\n{course}\n{instructor}\n{date}")
    format_paragraph(cover)

    doc.add_page_break()

    title_para = doc.add_paragraph(title)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(title_para)

    for para in text.split("\n"):
        p = doc.add_paragraph(para)
        format_paragraph(p)

# --- MLA ---
def apply_mla(doc, text, name, instructor, course, date, title):
    header = doc.add_paragraph(f"{name}\n{instructor}\n{course}\n{date}")
    format_paragraph(header)

    title_p = doc.add_paragraph(title)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(title_p)

    for para in text.split("\n"):
        p = doc.add_paragraph(para)
        format_paragraph(p)

# --- NEW SAVE FUNCTION ---
def save_file(text, file_path, filetype, style, name, instructor, course, date, title):

    # --- TXT ---
    if filetype == "TXT":
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(text)
        return

    # --- DOCX / PDF ---
    doc = Document()

    if style == "APA":
        apply_apa(doc, text, name, instructor, course, date, title)
    else:
        apply_mla(doc, text, name, instructor, course, date, title)

    # Save DOCX first
    docx_path = file_path.replace(".pdf", ".docx")
    doc.save(docx_path)

    # Convert to PDF if needed
    if filetype == "PDF":
        convert(docx_path, file_path)
        os.remove(docx_path)

# --- GENERATE ---
def generate():
    text = essay_box.get("1.0", tk.END).strip()
    filename = filename_entry.get()
    style = format_var.get()
    filetype = filetype_var.get()

    name = name_entry.get()
    instructor = instructor_entry.get()
    course = course_entry.get()
    date = date_entry.get()
    title = title_entry.get()

    if not text or not filename:
        messagebox.showerror("Error", "Essay and file name are required")
        return

    # FILE TYPE MAPPING
    ext_map = {
        "DOCX": ".docx",
        "PDF": ".pdf",
        "TXT": ".txt"
    }


    # --- CHOOSE SAVE LOCATION ---
    file_path = filedialog.asksaveasfilename(
        defaultextension=ext_map[filetype],
        initialfile=filename + ext_map[filetype],
        filetypes=[(f"{filetype} file", f"*{ext_map[filetype]}")]
    )

    if not file_path:
        return  # USER CANCELED

    processed = humanize_text(text)
    processed = additional_humanization(processed)

    save_file(processed, file_path, filetype, style, name, instructor, course, date, title)

    messagebox.showinfo("Success", f"Saved to:\n{file_path}")

# --- UI ---
root = tk.Tk()
root.title("AI Humanizer")

tk.Label(root, text="Paste Essay:").pack()
essay_box = tk.Text(root, height=15, width=80)
essay_box.pack()

tk.Label(root, text="File Name:").pack()
filename_entry = tk.Entry(root)
filename_entry.pack()

tk.Label(root, text="File Type:").pack()
filetype_var = tk.StringVar(value="DOCX")
ttk.Combobox(root, textvariable=filetype_var, values=["DOCX", "PDF", "TXT"]).pack()

format_var = tk.StringVar(value="APA")
ttk.Combobox(root, textvariable=format_var, values=["APA", "MLA"]).pack()

name_entry = tk.Entry(root)
name_entry.insert(0, "Your Name")
name_entry.pack()

instructor_entry = tk.Entry(root)
instructor_entry.insert(0, "Instructor")
instructor_entry.pack()

course_entry = tk.Entry(root)
course_entry.insert(0, "Course")
course_entry.pack()

date_entry = tk.Entry(root)
date_entry.insert(0, "Date")
date_entry.pack()

title_entry = tk.Entry(root)
title_entry.insert(0, "Essay Title")
title_entry.pack()

tk.Button(root, text="Generate Document", command=generate).pack(pady=10)

root.mainloop()